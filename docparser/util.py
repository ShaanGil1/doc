import math
import re

import mammoth
from docx import Document
from docx.oxml.ns import qn
from markdownify import MarkdownConverter

from .llm import build_payload, build_retry_message

from .models import (
    Sections_LLM,
    MatchedSection,
    MAX_RETRIES,
    MIN_SECTION_WORDS,
    MAX_SECTION_WORDS,
    MAX_DOCUMENT_WORDS,
)


# Document extraction

# Custom markdownify subclass: markdown has no <u>, so we keep it as inline HTML
class _MarkdownConverterWithUnderline(MarkdownConverter):
    def convert_u(self, el, text, parent_tags):
        return f'<u>{text}</u>'


def convert_to_markdown(docx_path):
    # Convert .docx to HTML with mammoth, then HTML to markdown with markdownify.
    # style_map='u => u' tells mammoth to keep underlines as <u> (default would strip).
    # apply_word_list_markers swaps the flat 1./2./3. markdownify produces with
    # the actual visible markers Word would have rendered (a., i., (1), etc.).
    with open(str(docx_path), 'rb') as docx_file:
        html = mammoth.convert_to_html(docx_file, style_map='u => u').value
    markdown = _MarkdownConverterWithUnderline(heading_style='ATX').convert(html)
    return apply_word_list_markers(markdown, Document(str(docx_path)))


def extract_paragraph_depths(docx_path):
    # Walk the .docx and return depth metadata for every non-empty paragraph.
    # Returns a list of {'text': str, 'depth': int or None}
    # The 'depth' value is the structural nesting level of the paragraph,
    # derived from .docx XML using ilvl and leading-tab count:

    document = Document(str(docx_path))
    rows = []

    # Top-level paragraphs
    for paragraph in document.paragraphs:
        row = paragraph_to_row(paragraph)
        if row is not None:
            rows.append(row)

    # For tables 
    def walk_table(table):
        for table_row in table.rows:
            seen_cells = []
            for cell in table_row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.append(cell._tc)
                for paragraph in cell.paragraphs:
                    row = paragraph_to_row(paragraph)
                    if row is not None:
                        rows.append(row)
                for nested_table in cell.tables:
                    walk_table(nested_table)

    for table in document.tables:
        walk_table(table)

    return rows

# Grab the text and depth for a paragraph, Trim to help LLM w/ overall struct 
def paragraph_to_row(paragraph):
    text = paragraph.text.strip()
    if not text:
        return None

    list_level = word_list_level(paragraph)
    tab_count = len(paragraph.text) - len(paragraph.text.lstrip('\t'))

    if list_level is None and tab_count == 0:
        depth = None
    else:
        depth = max(list_level or 0, tab_count)

    return {'text': text[:100], 'depth': depth}

# ilvl values
def word_list_level(paragraph):
    paragraph_props = paragraph._p.find(qn('w:pPr'))
    if paragraph_props is None:
        return None
    num_props = paragraph_props.find(qn('w:numPr'))
    if num_props is None:
        return None
    ilvl_element = num_props.find(qn('w:ilvl'))
    if ilvl_element is None:
        return None
    raw_value = ilvl_element.get(qn('w:val'))
    try:
        return int(raw_value) if raw_value is not None else None
    except ValueError:
        return None


# Word list-marker preservation
# Word's visible markers (a., i., (1)) are computed from numbering.xml at render
# time, not stored in the paragraph text. Mammoth+markdownify drop this and
# produce flat 1./2./3. The functions below read Word's numbering rules and
# substitute the correct markers back into the markdown.

LIST_LINE_PATTERN = re.compile(r'^(\s*)(\d+)\.\s+(.*)$')


# Replace markdownify's flat numbering with the actual Word markers
def apply_word_list_markers(markdown, doc):
    paragraph_markers = compute_list_markers(doc)
    if not paragraph_markers:
        return markdown

    lines = markdown.split('\n')
    cursor = 0
    for marker, paragraph_text in paragraph_markers:
        snippet = paragraph_text[:30].lower()
        if not snippet:
            continue
        # Find the next list line whose content matches this paragraph
        for i in range(cursor, len(lines)):
            match = LIST_LINE_PATTERN.match(lines[i])
            if not match:
                continue
            indent, _, content = match.groups()
            content_plain = re.sub(r'<[^>]+>|\*+|\[|\]\([^)]*\)', '', content).lower()
            if snippet[:15] in content_plain or content_plain[:15] in snippet:
                lines[i] = f'{indent}{marker} {content}'
                cursor = i + 1
                break
    return '\n'.join(lines)


# Walk paragraphs, return [(marker, text), ...] for every numbered (non-bullet) one
def compute_list_markers(doc):
    numbering_map = build_numbering_map(doc)
    if not numbering_map:
        return []

    counters = {}
    seen_restarts = set()  # (num_id, ilvl) pairs we've already restarted
    results = []
    for paragraph in doc.paragraphs:
        para_props = paragraph._p.find(qn('w:pPr'))
        if para_props is None:
            continue
        numbering_props = para_props.find(qn('w:numPr'))
        if numbering_props is None:
            continue
        num_id_el = numbering_props.find(qn('w:numId'))
        ilvl_el = numbering_props.find(qn('w:ilvl'))
        if num_id_el is None or ilvl_el is None:
            continue
        num_id = num_id_el.get(qn('w:val'))
        ilvl = int(ilvl_el.get(qn('w:val')))
        list_def = numbering_map.get(num_id)
        if not list_def:
            continue
        levels = list_def['levels']
        if ilvl not in levels:
            continue
        if levels[ilvl]['numFmt'] == 'bullet':
            continue

        # Honor explicit restart hints from numbering.xml: first time we see a
        # restart_level, force the counter back so it'll start fresh
        restart_key = (num_id, ilvl)
        if ilvl in list_def['restart_levels'] and restart_key not in seen_restarts:
            counters.pop((num_id, ilvl), None)
            seen_restarts.add(restart_key)

        # Reset deeper-level counters when this level advances
        for key in list(counters.keys()):
            if key[0] == num_id and key[1] > ilvl:
                del counters[key]

        # Increment this level's counter
        counter_key = (num_id, ilvl)
        counters[counter_key] = counters.get(counter_key, levels[ilvl]['start'] - 1) + 1

        # Render lvlText template, substituting %1, %2, ... with formatted counters
        marker = levels[ilvl]['lvlText']
        for level_index in range(ilvl + 1):
            if level_index not in levels:
                continue
            n = counters.get((num_id, level_index), levels[level_index]['start'])
            marker = marker.replace(f'%{level_index + 1}', format_list_number(n, levels[level_index]['numFmt']))

        results.append((marker, paragraph.text.strip()))

    return results


# Read numbering.xml -> {numId: {'levels': {ilvl: {numFmt, lvlText, start}}, 'restart_levels': set}}
def build_numbering_map(doc):
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return {}

    abstract_defs = {}
    for abstract_num in numbering_part.element.findall(qn('w:abstractNum')):
        abstract_num_id = abstract_num.get(qn('w:abstractNumId'))
        levels = {}
        for level_def in abstract_num.findall(qn('w:lvl')):
            ilvl = int(level_def.get(qn('w:ilvl')))
            num_fmt_el = level_def.find(qn('w:numFmt'))
            lvl_text_el = level_def.find(qn('w:lvlText'))
            start_el = level_def.find(qn('w:start'))
            levels[ilvl] = {
                'numFmt': num_fmt_el.get(qn('w:val')) if num_fmt_el is not None else 'decimal',
                'lvlText': lvl_text_el.get(qn('w:val')) if lvl_text_el is not None else f'%{ilvl + 1}.',
                'start': int(start_el.get(qn('w:val'))) if start_el is not None else 1,
            }
        abstract_defs[abstract_num_id] = levels

    result = {}
    for num in numbering_part.element.findall(qn('w:num')):
        num_id = num.get(qn('w:numId'))
        ref = num.find(qn('w:abstractNumId'))
        if ref is None:
            continue
        abstract_num_id = ref.get(qn('w:val'))
        levels = dict(abstract_defs.get(abstract_num_id, {}))

        # Look for lvlOverride entries that signal "restart at this level"
        restart_levels = set()
        for override in num.findall(qn('w:lvlOverride')):
            ilvl = int(override.get(qn('w:ilvl')))
            start_override = override.find(qn('w:startOverride'))
            if start_override is None:
                continue
            override_value = int(start_override.get(qn('w:val')))
            restart_levels.add(ilvl)
            if ilvl in levels:
                levels[ilvl] = dict(levels[ilvl])
                levels[ilvl]['start'] = override_value

        result[num_id] = {'levels': levels, 'restart_levels': restart_levels}

    return result


# Render an int per Word's various numbering formats
def format_list_number(n, num_fmt):
    if num_fmt == 'decimal':
        return str(n)
    if num_fmt == 'lowerLetter':
        return _to_lower_letter(n)
    if num_fmt == 'upperLetter':
        return _to_lower_letter(n).upper()
    if num_fmt == 'lowerRoman':
        return _to_lower_roman(n)
    if num_fmt == 'upperRoman':
        return _to_lower_roman(n).upper()
    return str(n)


# 1=a, 2=b, ..., 26=z, 27=aa, 28=ab
def _to_lower_letter(n):
    if n <= 0:
        return ''
    result = ''
    while n > 0:
        n -= 1
        result = chr(ord('a') + n % 26) + result
        n //= 26
    return result


# 1=i, 2=ii, 4=iv, 9=ix, ...
def _to_lower_roman(n):
    pairs = [(1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'),
             (100, 'c'), (90, 'xc'), (50, 'l'), (40, 'xl'),
             (10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i')]
    result = ''
    for value, symbol in pairs:
        while n >= value:
            result += symbol
            n -= value
    return result


# LLM call + line matching

# Call the LLM to get section boundaries, with retry on no-match.
def request_sections(llm, payload, max_retries=MAX_RETRIES):

    structured_llm = llm.with_structured_output(Sections_LLM)
    user_message = payload['user_message']
    lines = payload['markdown_lines']
    best_attempt_sections = None
    best_attempt_match_count = 0

    for attempt in range(max_retries):
        try:
            response = structured_llm.invoke([
                ('system', payload['system']),
                ('user', user_message),
            ])
        except Exception:
            continue

        if response is None or not response.sections:
            continue

        all_sections = match_sections(response.sections, lines)
        match_count = sum(1 for s in all_sections if s.matched and s.match_text)

        if match_count > best_attempt_match_count:
            best_attempt_sections = all_sections
            best_attempt_match_count = match_count

        if match_count > 0:
            return all_sections

        user_message = build_retry_message(
            payload['user_message'], all_sections, attempt, max_retries,
        )

    if best_attempt_match_count > 0:
        return best_attempt_sections
    return []


def match_sections(ai_sections, lines):
    # Match the sections
    sections = []
    cursor = 0

    for item in ai_sections:
        match_text = item.section_match_text
        if not match_text:
            continue

        line_index = find_matching_line(match_text, lines, search_from=cursor)
        if line_index is None:
            sections.append(MatchedSection(item.title, -1, match_text, matched=False))
        else:
            sections.append(MatchedSection(item.title, line_index + 1, match_text, matched=True))
            cursor = line_index + 1

    matched_only = [s for s in sections if s.matched]
    # Some edge casing likely never triggers
    if not matched_only or matched_only[0].start_line > 1:
        sections.insert(0, MatchedSection('Introduction', 1, '', matched=True))
    return sections


def find_matching_line(match_text, lines, search_from=0):
    # Locate match_text in lines using a 3-stage fallback.

    stripped_match = match_text.strip()
    if not stripped_match:
        return None

    # exact line match
    for i in range(search_from, len(lines)):
        if lines[i].strip() == stripped_match:
            return i

    # substring match (handles LLM dropping '## ' prefix etc.).
    if len(stripped_match) >= 8:
        for i in range(search_from, len(lines)):
            stripped_line = lines[i].strip()
            if len(stripped_line) < 8:
                continue
            if stripped_match in lines[i] or stripped_line in stripped_match:
                return i

    # normalized match (whitespace/punctuation/case differences)
    normalized_match = re.sub(r'[^a-z0-9]', '', stripped_match.lower())
    if len(normalized_match) >= 4:
        for i in range(search_from, len(lines)):
            normalized_line = re.sub(r'[^a-z0-9]', '', lines[i].lower())
            if normalized_line == normalized_match:
                return i
    return None

# Post-processing:
def section_is_mostly_table(content):
    # Heuristic: if >= 70% markdown table and we should avoid splitting it

    table_words = 0
    total_words = 0
    for line in content.split('\n'):
        word_count = len(line.split())
        total_words += word_count
        if line.strip().startswith('|'):
            table_words += word_count
    if total_words == 0:
        return False
    return table_words / total_words >= 0.7


def merge_and_split_sections(matched_sections, lines, min_words, max_words):
    # Merge sections smaller than min_words into neighbors, then split sections larger than max_words at sentence boundaries
    sections = sorted(
        [s for s in matched_sections if s.matched],
        key=lambda s: s.start_line,
    )
    if not sections:
        return []

    # Helper
    def words_at(i):
        start = sections[i].start_line - 1
        end = sections[i + 1].start_line - 1 if i + 1 < len(sections) else len(lines)
        return sum(len(line.split()) for line in lines[start:end])

    # Merge undersized sections
    if len(sections) > 1:
        i = 0
        while i < len(sections):
            if words_at(i) >= min_words:
                i += 1
                continue
            if i == 0 and len(sections) > 1:
                # First section too small - extend the next section back to line 1
                next_section = sections[1]
                sections[1] = MatchedSection(
                    title=next_section.title,
                    start_line=sections[0].start_line,
                    match_text=next_section.match_text,
                    matched=True,
                )
                sections.pop(0)
            elif i > 0:
                sections.pop(i)  # fold this section back into the previous one
            else:
                i += 1

    # Split oversized at sentence boundaries
    output = []
    for index, section in enumerate(sections):
        start = section.start_line - 1
        end = sections[index + 1].start_line - 1 if index + 1 < len(sections) else len(lines)
        content = '\n'.join(lines[start:end]).strip()

        # Table-dominant sections are kept whole regardless of word count
        if section_is_mostly_table(content):
            output.append({'title': section.title, 'content': content})
            continue

        sentence_seperator = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')

        sentences = [s for s in sentence_seperator.split(content) if s.strip()]
        total_words = sum(len(s.split()) for s in sentences)

        if total_words <= max_words or len(sentences) < 2:
            output.append({'title': section.title, 'content': content})
            continue

        # Pick split points at sentence boundaries closest to even word distribution
        num_parts = math.ceil(total_words / max_words)
        target = total_words / num_parts
        cumulative = [0]
        for s in sentences:
            cumulative.append(cumulative[-1] + len(s.split()))

        split_points = [0]
        for part_index in range(1, num_parts):
            ideal = part_index * target
            previous = split_points[-1]
            best = previous + 1
            best_diff = abs(cumulative[best] - ideal)
            for candidate in range(previous + 2, len(sentences)):
                diff = abs(cumulative[candidate] - ideal)
                if diff < best_diff:
                    best_diff = diff
                    best = candidate
            split_points.append(best)
        split_points.append(len(sentences))

        for n in range(len(split_points) - 1):
            piece = ' '.join(sentences[split_points[n]:split_points[n + 1]])
            output.append({
                'title': f'{section.title} ({n + 1}/{num_parts})',
                'content': piece,
            })

    return output

# Main pipeline
def run_sectioning(
    llm,
    docx_path,
    max_retries=MAX_RETRIES,
    max_words=MAX_DOCUMENT_WORDS,
    min_section_words=None,
    max_section_words=None,
):
    # Convert a .docx into editable sections. The whole pipeline.
    # 1. convert_to_markdown() turns the .docx into markdown text.
    # 2. build_payload() converts the docx, extracts metadata, assembles the LLM user message.
    # 3. request_sections() calls the LLM with retries on no-match.
    # 4. merge_and_split_sections() merges undersized sections into neighbors and splits oversized

    markdown = convert_to_markdown(str(docx_path))
    metadata_rows = extract_paragraph_depths(docx_path)
    payload = build_payload(markdown, metadata_rows, max_words=max_words)
    matched_sections = request_sections(llm, payload, max_retries=max_retries)

    if not matched_sections:
        return [{'title': 'Document', 'content': '\n'.join(payload['markdown_lines'])}]

    # Adaptive thresholds: shrink for short docs so legitimate sections
    adaptive_min, adaptive_max = adaptive_thresholds(payload['word_count'])
    final_min = min_section_words if min_section_words is not None else adaptive_min
    final_max = max_section_words if max_section_words is not None else adaptive_max

    return merge_and_split_sections(
        matched_sections,
        payload['markdown_lines'],
        min_words=final_min,
        max_words=final_max,
    )

# Change if params feel off
def adaptive_thresholds(total_words):
    min_words = min(MIN_SECTION_WORDS, max(20, total_words // 15))
    max_words = min(MAX_SECTION_WORDS, max(150, total_words // 5))
    return min_words, max_words
