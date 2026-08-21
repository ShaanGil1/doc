from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Dict, List, NamedTuple, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt

import config as cfg
import helpers
import logic

# RENDERING
# Formatting lives in Word styles, created once by create_styles(). 
def create_styles(document):
    """Turn every entry in config.STYLES into a real Word style"""
    for spec in cfg.STYLES.values():
        helpers.define_style(
            document, spec.name, based_on=spec.base, font=spec.font,
            size_pt=spec.size, color=spec.color, bold=spec.bold,
            italic=spec.italic, all_caps=spec.caps,
            align=spec.align, space_before_pt=spec.before,
            space_after_pt=spec.after)
    # Runs that need underlining inside an otherwise plain paragraph, which is every lead-in title
    helpers.define_style(document, cfg.UNDERLINE_STYLE, character=True, underline=True)

def write(document, key: str, text: str = "", level: int = 0, lead_in: str = "", num_id: int = 0):
    """One paragraph in the named style"""
    spec = cfg.STYLES[key]
    paragraph = document.add_paragraph(style=spec.name)
    paragraph_format = paragraph.paragraph_format

    if spec.first_line:
        paragraph_format.first_line_indent = Inches(spec.first_line)
    if spec.indent:
        paragraph_format.left_indent = Inches(cfg.INDENT_STEP_IN * (level + 1))
        paragraph_format.first_line_indent = Inches(-cfg.INDENT_STEP_IN)
    if spec.page_break:
        paragraph_format.page_break_before = True

    if lead_in:
        run = paragraph.add_run(lead_in + ":")
        run.style = document.styles[cfg.UNDERLINE_STYLE]
        if text:
            paragraph.add_run(" ")
    if text:
        run = paragraph.add_run(text + ("" if lead_in else spec.suffix))
        # Underline goes on the RUN, not the paragraph style
        if spec.underline and not lead_in:
            run.style = document.styles[cfg.UNDERLINE_STYLE]

    if spec.numbered and num_id:
        helpers.Numbering.apply_to_paragraph(paragraph, num_id, level)
    return paragraph

def write_image(document, path, width_in, height_in, key="cover_line"):
    paragraph = document.add_paragraph(style=cfg.STYLES[key].name)
    paragraph.alignment = cfg.LEFT
    paragraph.add_run().add_picture(path, width=Inches(width_in), height=Inches(height_in))
    return paragraph


def new_cascade_list(numbering, cascade=None) -> int:
    """A fresh numId each call, which is what makes numbering restart"""
    return numbering.create_list(True, cascade=cascade or cfg.CASCADE, suffix=cfg.NUMBER_SUFFIX)

# READING THE MARKDOWN
MARKUP = re.compile(r"[*_`~#]+")
TRAILING_PUNCTUATION = re.compile(r"[\s:.\-–—;,]+$")
WHITESPACE = re.compile(r"\s+")

def normalize(text: str) -> str:
    """'Summary of  Changes :' and '**SUMMARY OF CHANGES**' both become 'SUMMARY OF CHANGES'"""
    text = MARKUP.sub("", unicodedata.normalize("NFKC", text or ""))
    return TRAILING_PUNCTUATION.sub("", WHITESPACE.sub(" ", text).strip()).upper()

def heading_text(node) -> str:
    return logic.spans_text(logic.inline_spans(logic.find_inline_child(node)))

def heading_level(node) -> int:
    return int(node.tag[1:] or 1)

def split_lead_in(text: str) -> Tuple[str, str]:
    """'Scope: rest' -> ('Scope', 'rest'). A colon is the only marker, so
    'Scope. rest' is an ordinary sentence and comes back as ('', text)"""
    position = text.find(":") if text else -1
    if position < 0:
        return "", text
    head, rest = text[:position], text[position + 1:].strip()
    if (not head.strip() or not rest
            or len(head) > cfg.LEAD_IN["max_chars"]
            or len(head.split()) > cfg.LEAD_IN["max_words"]
            or any(char in head for char in cfg.LEAD_IN["stop_chars"])):
        return "", text
    return head.strip(), rest

SECTION_ALIASES = {normalize(alias): name for name, extra in cfg.REQUIRED_SECTIONS.items() for alias in (name,) + tuple(extra)}

def split_sections(markdown_text: str) -> Tuple[Dict[str, list], List[str]]:
    """Slice sections_input into the required slots"""
    slots: Dict[str, list] = {name: [] for name in cfg.REQUIRED_SECTIONS}
    found, findings = set(), []
    if not (markdown_text or "").strip():
        return slots, ["MISSING required section: %s" % name for name in slots]

    current, current_level = None, 99
    for node in logic.parse_markdown(markdown_text).children:
        if node.type == "heading":
            level, label = heading_level(node), heading_text(node)
            name = SECTION_ALIASES.get(normalize(label))
            if level <= current_level:
                if name and name not in found:
                    found.add(name)
                    current, current_level = name, level
                    continue
                if name:
                    findings.append("DUPLICATE heading for section: %s" % name)
                else:
                    findings.append("UNMATCHED heading (kept): %s" % label)
                    if current is None:
                        continue
        if current is not None:
            slots[current].append(node)

    missing = ["MISSING required section: %s" % name for name in slots if name not in found]
    return slots, missing + findings

class TocEntry(NamedTuple):
    level: int
    text: str
    bookmark: str

class Outline(NamedTuple):
    groups: list        # [(title node or None, [body nodes])] per enclosure
    offset: int         # subtract from a heading level to normalise it
    entries: list       # every TocEntry, in document order
    titles: list        # enclosure titles, for the signature block


def outline(markdown_text: str) -> Outline:
    """Everything the later pages need to know, worked out before writing
    Whatever the author used as their top heading level becomes the enclosure
    level, so an H1 led document and an H2 led one give the same structure"""
    tree = logic.parse_markdown(markdown_text or "")
    levels = [heading_level(n) for n in tree.children if n.type == "heading"]
    offset = (min(levels) if levels else 1) - 1

    groups, title, body = [], None, []
    for node in tree.children:
        if node.type == "heading" and heading_level(node) - offset == 1:
            if title is not None or body:
                groups.append((title, body))
            title, body = node, []
        else:
            body.append(node)
    if title is not None or body:
        groups.append((title, body))

    entries, titles = [], []

    def add(level, text):
        entries.append(TocEntry(level, text, "DLAIREF%d" % (len(entries) + 1)))

    for index, (group_title, group_body) in enumerate(groups, start=1):
        name = heading_text(group_title) if group_title else "ENCLOSURE"
        titles.append(name)
        add(1, cfg.BACK["encl_display"] % (index, name))
        for node in group_body:
            if node.type == "heading":
                level = max(1, heading_level(node) - offset)
                if level <= cfg.BACK["toc_depth"]:
                    add(level, heading_text(node))

    # the back matter always exists, so its rows are fixed
    add(1, cfg.BACK["encl_last"])
    for part in cfg.BACK["glossary_parts"]:
        add(2, part)
    for name in cfg.BACK["trailing_lists"]:
        add(1, name)

    return Outline(groups, offset, entries, titles)

# WRITING THE MARKDOWN
class DlaiWriter(logic.DocxWriter):
    """logic.DocxWriter with the template's rules applied"""
    def __init__(self, document, opts, numbering, num_id, offset=0,
                 base_level=-1, min_level=0,
                 blocks=("enclosure", "encl_h2", "encl_h3"), bookmarks=None,
                 body_key="prose", lead_key="subsection"):
        super().__init__(document, opts)
        self.numbering = numbering      # shared
        self.num_id = num_id
        self.offset = offset
        self.min_level = min_level
        # Numbering level of the heading currently in scope
        self.level = base_level
        self.blocks = blocks
        self.bookmarks = bookmarks if bookmarks is not None else []
        # DEFINITIONS overrides these so its entries sit left 
        self.body_key, self.lead_key = body_key, lead_key

    def content_level(self, position) -> int:
        return self.level + 1 + max(0, position.list_depth)

    def write_heading(self, node, position):
        level = max(1, heading_level(node) - self.offset)
        key = self.blocks[min(level, len(self.blocks)) - 1]
        num_level = max(self.min_level, level - 1)
        paragraph = write(self.document, key, heading_text(node), level=num_level, num_id=self.num_id)
        if cfg.STYLES[key].numbered:
            self.level = num_level
        if self.bookmarks and level <= cfg.BACK["toc_depth"]:
            helpers.add_bookmark(paragraph, self.bookmarks.pop(0))
        self.last_paragraph = paragraph
        return paragraph

    def write_paragraph(self, node, position):
        text = logic.spans_text(
            logic.inline_spans(logic.find_inline_child(node))).strip()
        if not text:
            return None
        lead_in, rest = split_lead_in(text)
        self.last_paragraph = write(
            self.document, self.lead_key if lead_in else self.body_key, rest,
            level=self.content_level(position), lead_in=lead_in,
            num_id=self.num_id)
        return self.last_paragraph

    def new_paragraph(self, position, style: Optional[str] = None):
        level = self.content_level(position)
        paragraph = write(self.document, self.body_key, level=level)
        if position.list_depth >= 0:
            helpers.Numbering.apply_to_paragraph(paragraph, self.num_id, level)
        self.last_paragraph = paragraph
        return paragraph



# 4. HARDCODED TEMPLATE
def setup_page(document, doc_title: str):
    """Page geometry, the Normal style, then the header and footer"""
    for section in document.sections:
        section.page_width = Inches(cfg.PAGE["width_in"])
        section.page_height = Inches(cfg.PAGE["height_in"])
        section.top_margin = section.bottom_margin = Inches(cfg.PAGE["margin_in"])
        section.left_margin = section.right_margin = Inches(cfg.PAGE["margin_in"])

    helpers.define_style(document, "Normal", font=cfg.BODY_FONT,
                         size_pt=cfg.BODY_SIZE, color=cfg.BODY_COLOR,
                         space_after_pt=cfg.GAP)
    create_styles(document)

    section = document.sections[0]
    for area, align, text in ((section.header, cfg.PAGE["header_align"], doc_title),
                              (section.footer, cfg.PAGE["footer_align"], None)):
        area.is_linked_to_previous = False
        paragraph = area.paragraphs[0]
        paragraph.alignment = align
        if text:
            paragraph.add_run(text)
        else:
            helpers.add_page_number_field(paragraph)


def write_cover(document, numbering, doc_title: str):
    cover = cfg.COVER
    for path, _, _ in (cover["seal"], cover["rule"]):
        if not Path(path).is_file():
            raise FileNotFoundError(
                "Cover image missing: %s\nThe assets folder must sit next to "
                "config.py." % path)
    write_image(document, *cover["seal"])
    write(document, "agency", cover["agency_name"])
    write(document, "doc_type", cover["doc_type"])
    write(document, "cover_line", doc_title)
    write(document, "cover_line", cover["effective_text"])
    write_image(document, *cover["rule"], key="agency")

    label = "cover_label"
    for index, text in enumerate(cover["labels"]):
        paragraph = write(document, label, text)
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(
                cover["label_space_before_pt"])

    num_id = new_cascade_list(numbering, cover["ref_cascade"])
    reference = "cover_ref"
    for index in range(1, cover["ref_count"] + 1):
        paragraph = write(document, reference, cover["ref_pattern"] % index,
                          num_id=num_id)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(cover["ref_indent_in"])
        paragraph_format.first_line_indent = Inches(-cfg.INDENT_STEP_IN)
        if index == cover["ref_count"]:
            paragraph_format.space_after = Pt(cfg.GAP)


def write_signature_block(document, titles: List[str]):
    """References is always Enclosure 1"""
    back = cfg.BACK
    titles = [back["encl_first"]] + [t for t in titles if t != back["encl_first"]]
    for line in back["sig_lines"]:
        write(document, "signature", line)
    write(document, "encl_label", back["encl_label"])
    for index, title in enumerate(titles, start=1):
        write(document, "encl_item",
              back["encl_pattern"] % (index, title))
    glossary = write(document, "encl_item", back["encl_last"])
    glossary.paragraph_format.first_line_indent = Inches(0)


def write_table_of_contents(document, entries: List[TocEntry]):
    """The shape is fixed; the rows come from the real document's outline"""
    write(document, "toc_title", cfg.BACK["toc_title"])

    rows = {1: "toc_1", 2: "toc_2", 3: "toc_3"}
    for entry in entries:
        key = rows.get(entry.level, "toc_3")
        indent = cfg.BACK["toc_indent_in"] * (entry.level - 1)
        paragraph = document.add_paragraph(style=cfg.STYLES[key].name)
        paragraph.paragraph_format.left_indent = Inches(indent)
        helpers.add_dot_leader_tab(paragraph, cfg.BACK["toc_tab_in"] - indent)
        paragraph.add_run(entry.text + "\t")
        helpers.add_page_reference_field(paragraph, entry.bookmark)
        helpers.wrap_paragraph_in_internal_link(paragraph, entry.bookmark)


def write_back_matter(document, bookmarks: List[str]):
    """Glossary with its parts, then the tables and figures lists"""
    back = cfg.BACK
    paragraph = write(document, "enclosure", back["encl_last"])
    helpers.add_bookmark(paragraph, bookmarks.pop(0))
    for part in back["glossary_parts"]:
        paragraph = write(document, "glossary_part", part)
        helpers.add_bookmark(paragraph, bookmarks.pop(0))
        write(document, "prose", "[Entries come from markdown.]")
    for name in back["trailing_lists"]:
        paragraph = write(document, "enclosure", name)
        helpers.add_bookmark(paragraph, bookmarks.pop(0))
        write(document, "prose", "[List generated from captions.]")

# PUTTING IT TOGETHER
def build_dlai(markdown_input: str = "", sections_input: str = "",
              doc_title: str = "PLACEHOLDER TITLE") -> Tuple[bytes, List[str]]:
    opts = helpers.DocxOptions(body_font=cfg.BODY_FONT,
                               body_size_pt=cfg.BODY_SIZE)
    document = Document()
    numbering = helpers.Numbering(document)

    plan = outline(markdown_input)
    bookmarks = [entry.bookmark for entry in plan.entries]

    setup_page(document, doc_title)
    write_cover(document, numbering, doc_title)
    findings = write_required_sections(document, opts, numbering, sections_input)
    write_signature_block(document, plan.titles)
    write_table_of_contents(document, plan.entries)
    write_enclosures(document, opts, numbering, plan, bookmarks)
    write_back_matter(document, bookmarks)

    helpers.update_fields_on_open(document)
    return helpers.document_to_bytes(document), findings


def write_required_sections(document, opts, numbering, sections_input: str) -> List[str]:
    """One shared numId so the sections count 1..N instead of each restarting"""
    slots, findings = split_sections(sections_input)
    num_id = new_cascade_list(numbering)

    for index, (name, nodes) in enumerate(slots.items()):
        paragraph = write(document, "section", name, num_id=num_id)
        if index == 0 and cfg.PAGE["break_after_cover"]:
            paragraph.paragraph_format.page_break_before = True
        if not nodes:
            write(document, "prose", cfg.MISSING_PLACEHOLDER,
                  level=1, num_id=num_id)
            continue
        key = "definition" if name in cfg.FLAT_SECTIONS else None
        DlaiWriter(document, opts, numbering, num_id, base_level=0,
                  min_level=1, blocks=("subsection",),
                  body_key=key or "prose", lead_key=key or "subsection"
                  ).write_blocks(nodes, logic.Position())
    return findings


def write_enclosures(document, opts, numbering, plan, bookmarks: List[str]):
    """A fresh numId per enclosure, which restarts its numbering at 1"""
    for index, (title, body) in enumerate(plan.groups, start=1):
        num_id = new_cascade_list(numbering)
        label = cfg.BACK["encl_display"] % (
            index, heading_text(title) if title else "ENCLOSURE")
        heading = write(document, "enclosure", label)
        helpers.add_bookmark(heading, bookmarks.pop(0))
        DlaiWriter(document, opts, numbering, num_id, offset=plan.offset,
                  bookmarks=bookmarks).write_blocks(body, logic.Position())
