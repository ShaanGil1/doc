"""python md_to_docx/tests/check.py

Proves the pipeline still meets the locked standard on the inputs in tests/:
findings as expected, validator clean, no code block ever produced from list
input, and numbering unchanged under random re-indentation."""

import random, re, sys, tempfile
from pathlib import Path

sys.path[:0] = [str(Path(__file__).resolve().parent.parent), str(Path(__file__).resolve().parent)]
import template_processor, boundaries, config, validate, logic
from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
SCRATCH = Path(tempfile.gettempdir()) / "md_to_docx_check.docx"
results = []


def check(name, ok, detail=""):
    results.append((name, ok, detail))


def build(text, tmpl="SOP"):
    b, f = template_processor.build_dlai(text, "", "CHECK", tmpl)
    SCRATCH.write_bytes(b)
    problems, _ = validate.validate(str(SCRATCH))
    return f, problems


def levels():
    out = []
    for p in Document(str(SCRATCH)).paragraphs:
        numpr = p._p.pPr.find(qn("w:numPr")) if p._p.pPr is not None else None
        if numpr is not None and p.text.strip():
            out.append((numpr.find(qn("w:ilvl")).get(qn("w:val")), p.text[:30]))
    return out


sop = (HERE / "sop_template.md").read_text()
f, problems = build(sop)
check("SOP validates", not problems, str(problems[:1]))
check(
    "SOP findings",
    {x.split("(")[0].strip() for x in f}
    >= {
        "bold mode: no # headings, titles identified from bold lines",
        "SIGNATURE BLOCK taken",
        "TABLE OF CONTENTS block dropped; the table is generated",
        "MISSING required section: POLICY",
    },
    str(f),
)
ref = levels()

f, problems = build((HERE / "demo.md").read_text(), None)
check("demo.md (# path) validates with 3 MISSING", not problems and sum("MISSING" in x for x in f) == 3, str(f))

for name in (
    "references_torn.md",
    "references_clean.md",
    "list_examples.md",
    "glossary_both_parts.md",
    "deep_nesting.md",
):
    f, problems = build((HERE / name).read_text())
    check("%s validates" % name, not problems, str(problems[:1]))

MARK = re.compile(r"^(\s*)((?:\d+\.|[a-z]\.|\(\d+\)|\([a-z]\))\s)", re.M)


def fuzz(text, seed):
    rnd = random.Random(seed)
    lines = MARK.sub(lambda m: rnd.choice(["", "  ", "    ", "\t", "        "]) + m.group(2), text).splitlines()
    out = []
    for i, line in enumerate(lines):
        out.append(line)
        nxt = lines[i + 1] if i + 1 < len(lines) else ""
        if MARK.match(line) and MARK.match(nxt) and rnd.random() < 0.5:
            out.append("")
    return "\n".join(out)


stable = True
for seed in range(10):
    build(fuzz(sop, seed))
    stable &= levels() == ref
check("10 re-indented SOP variants give identical numbering", stable)


def no_code(text):
    doc = boundaries.build_document(text)
    blocks = [b for b in doc.sections.values() if b] + doc.enclosures
    return all(
        n.type not in ("code_block", "fence")
        for b in blocks
        for n in logic.parse_markdown("\n".join(template_processor.normalize_lists(b.lines))).children
    )


check(
    "no code block from any list-bearing input",
    all(no_code(fuzz(sop, s)) for s in range(10)) and no_code((HERE / "references_torn.md").read_text()),
)

# edge cases: titles and structure
f, problems = build((HERE / "edge_titles.md").read_text(encoding="utf-8"))
doc = boundaries.build_document((HERE / "edge_titles.md").read_text(encoding="utf-8"))
check("edge_titles validates", not problems, str(problems[:1]))
check("BOM and __bold__: OPR read from '__OPR:__' line", doc.cover.get("opr") == "J6 Logistics", str(doc.cover))
check(
    "cover references lose their a./b. markers",
    doc.cover.get("references", "").startswith("DoD Instruction"),
    doc.cover.get("references", "")[:40],
)
check(
    "inline-body titles: colon outside, inside, and with a period",
    all(doc.sections[n] is not None for n in ("DEFINITIONS", "INFORMATION REQUIREMENTS", "RELEASABILITY", "PURPOSE")),
)
check("inline body becomes the first body line", doc.sections["DEFINITIONS"].lines[0] == "See Glossary.")
check("empty section reported and placeholder written", any("EMPTY section: SUMMARY OF CHANGES" in x for x in f))
check("duplicate section reported once", sum("DUPLICATE heading for section: POLICY" in x for x in f) == 1)
check(
    "enclosure title variants (dash, period, trailing colon)",
    [b.title for b in doc.enclosures] == ["References", "Responsibilities", "Procedures"],
    str([b.title for b in doc.enclosures]),
)
check(
    "glossary parts with arabic numerals, tab-separated row kept",
    doc.glossary.abbreviations is not None and any("\t" in l for l in doc.glossary.abbreviations.lines),
)
check(
    "signature from fence: 2 lines",
    doc.signature == ["JANE Q. DOE", "Director, Information Operations"],
    str(doc.signature),
)
# edge cases: bodies
f, problems = build((HERE / "edge_bodies.md").read_text())
doc = boundaries.build_document((HERE / "edge_bodies.md").read_text())
check("edge_bodies validates", not problems, str(problems[:1]))
check("empty and whitespace-only enclosures reported", sum("has no content" in x for x in f) == 2)
check("torn '(' left as written (a paragraph, then the a. item)", doc.sections["PROCEDURES"].lines[0] == "(")
lv = levels()
check(
    "(1)-first section renders flat at a.",
    [l for l, t in lv if t.startswith(("A section whose", "Second item."))] == ["1", "1"],
)
check("shape repeat under (a) reads as (1)", [l for l, t in lv if t.startswith("Deepest")] == ["2"])
# the signature written as a plain marker line, followed by a fenced Enclosure(s) list (a real frontend shape)
f, problems = build((HERE / "signature_after_lines.md").read_text())
doc = boundaries.build_document((HERE / "signature_after_lines.md").read_text())
check("plain-marker signature with a fenced Enclosure(s) list after it: builds, no code block error", not problems)
check("  signature is the three lines between the marker and the list", doc.signature == ["[INPUT REQUIRED: Printed Name of Authorizing Official]", "Director, DLA Finance (J8)", "Defense Logistics Agency"], str(doc.signature))
check("  authored list dropped with a finding, ToC dropped, 5 enclosures", any("Enclosure(s) list under the signature dropped (6" in x for x in f) and any("TABLE OF CONTENTS block dropped" in x for x in f) and len(doc.enclosures) == 5)
check("  nothing of the signature region leaks into EXPIRATION DATE", not any("Enclosure" in l or "SIGNATURE" in l for l in doc.sections["EXPIRATION DATE"].lines))

# optional sections: absent -> left out entirely, numbering closes up, different finding
saved_env = config.SECTIONS["POLICY"]
config.SECTIONS["POLICY"] = config.Section(optional=True)
f, problems = build(sop)
heads = [p.text for p in Document(str(SCRATCH)).paragraphs if p.style.name == "DLAI Section"]
check(
    "optional POLICY absent: no heading, no placeholder, finding says optional",
    "POLICY:" not in heads
    and len(heads) == 10
    and not any("section not found" in p.text for p in Document(str(SCRATCH)).paragraphs)
    and any("optional section not present: POLICY" in x for x in f),
)
f, problems = build((HERE / "example_dlai_finance.md").read_text())
check(
    "optional POLICY present: prints exactly like a required section",
    "POLICY:" in [p.text for p in Document(str(SCRATCH)).paragraphs if p.style.name == "DLAI Section"],
)
config.SECTIONS["POLICY"] = saved_env

f, problems = build("1. **PURPOSE:**\nfine\n\n```\nnot a signature\n```\n")
check("a fence in a body never fails: markers stripped, lines print as text, finding says so", not problems and any("fence markers stripped in 'PURPOSE'" in x for x in f))
f, problems = build("1. **PURPOSE:**\nx\n**Enclosure 1: A**\na\n")
check("no signature block: nothing printed, finding says so", not problems and any("SIGNATURE BLOCK not found" in x for x in f))

w = max(len(n) for n, _, _ in results)
for name, ok, d in results:
    print("%s  %s%s" % ("PASS" if ok else "FAIL", name.ljust(w), "   " + d if not ok and d else ""))
print("\n%d/%d" % (sum(ok for _, ok, _ in results), len(results)))
sys.exit(0 if all(ok for _, ok, _ in results) else 1)
