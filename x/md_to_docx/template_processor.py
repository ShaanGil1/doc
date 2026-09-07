"""Renders a DlaiDocument as a DLA .docx: md_to_docx(document, title, template) -> (bytes, findings).
Reading is boundaries.build_document; formatting values live in config; helpers/logic are the generic layers."""

from __future__ import annotations

from pathlib import Path
from typing import List, NamedTuple, Optional, Tuple

from docx import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt

import boundaries
import config as cfg
import helpers
import logic
from rules import (
    FENCE_LINE,
    GLOSSARY_ENTRY,
    HEADING_LINE,
    LIST_MARKER,
    MARKER_DEPTH,
    RULE_LINE,
    SECTION_ALIASES,
    Block,
    ConversionError,
    DlaiDocument,
    normalize,
)


# =========================================================================== #
# 1. RENDERING: formatting lives in Word styles; a paragraph carries only what a style cannot (number, indent, break)
# =========================================================================== #
def create_styles(document):
    """Turn every entry in config.STYLES into a real Word style"""
    for spec in cfg.STYLES.values():
        helpers.define_style(
            document,
            spec.name,
            based_on=spec.base,
            font=spec.font,
            size_pt=spec.size,
            color=spec.color,
            bold=spec.bold,
            italic=spec.italic,
            all_caps=spec.caps,
            align=spec.align,
            space_before_pt=spec.before,
            space_after_pt=spec.after,
            line_spacing=spec.line,
            keep_with_next=False,
            keep_together=False,
        )
    # Runs that need underlining inside an otherwise plain paragraph, which is
    # every lead-in title. A character style keeps it reviewer editable
    helpers.define_style(document, cfg.UNDERLINE_STYLE, character=True, underline=True)


def write(
    document, key: str, text: str = "", level: int = 0, lead_in: str = "", num_id: int = 0, indent_in: float = None
):
    """One paragraph in the named style; lead_in is an underlined title opening it, indent_in overrides the indent
    step."""
    spec = cfg.STYLES[key]
    step = cfg.INDENT_STEP_IN if indent_in is None else indent_in
    if spec.page_break:
        new_page(document)
    paragraph = document.add_paragraph(style=spec.name)
    paragraph_format = paragraph.paragraph_format

    if spec.first_line:
        paragraph_format.first_line_indent = Inches(spec.first_line)
    if spec.indent:
        paragraph_format.left_indent = Inches(step * (level + 1))
        paragraph_format.first_line_indent = Inches(-step)

    if lead_in:
        run = paragraph.add_run(lead_in + spec.lead_suffix)
        run.style = document.styles[cfg.UNDERLINE_STYLE]
        if text:
            paragraph.add_run(spec.lead_gap)
    if text:
        run = paragraph.add_run(text + ("" if lead_in else spec.suffix))
        # underline on the run, not the style, or the auto-number would be underlined too
        if spec.underline and not lead_in:
            run.style = document.styles[cfg.UNDERLINE_STYLE]

    if spec.numbered and num_id:
        helpers.Numbering.apply_to_paragraph(paragraph, num_id, level)
    return paragraph


def new_page(document):
    """Next-page section break (as the template does). A page-break-before property would show Word's black square."""
    document.add_section(WD_SECTION.NEW_PAGE)


def write_image(document, path, width_in, height_in, key="cover_line"):
    paragraph = document.add_paragraph(style=cfg.STYLES[key].name)
    paragraph.alignment = cfg.LEFT
    paragraph.add_run().add_picture(path, width=Inches(width_in), height=Inches(height_in))
    return paragraph


def new_cascade_list(numbering, cascade=None) -> int:
    """A fresh numId each call, which is what makes numbering restart"""
    return numbering.create_list(
        True, cascade=cascade or cfg.CASCADE, suffix=cfg.NUMBER_SUFFIX, underline_repeat=cfg.CASCADE_REPEAT_UNDERLINE
    )


# =========================================================================== #
# 2. BODIES  -  one block at a time
# =========================================================================== #
def heading_text(node) -> str:
    return logic.spans_text(logic.inline_spans(logic.find_inline_child(node)))


def heading_level(node) -> int:
    return int(node.tag[1:] or 1)


def paragraph_lines(node) -> List[str]:
    """The paragraph's source lines, split at soft and hard breaks, inner spaces intact (a two-column glossary needs
    them)."""
    inline = logic.find_inline_child(node)
    lines, current = [], []
    for child in (inline.children if inline is not None else []):
        if child.type in ("softbreak", "hardbreak"):
            lines.append("".join(current))
            current = []
        elif child.children:
            spans: list = []
            logic.collect_spans(child, {}, spans)
            current.append(logic.spans_text(spans))
        else:
            current.append(child.content or "")
    lines.append("".join(current))
    return [line.strip() for line in lines if line.strip()]


def glossary_entry(line: str):
    """(term, meaning) for a glossary row, else None. The term ends at the first tab, run of 2+ spaces, or colon."""
    match = GLOSSARY_ENTRY.match(line)
    return (match.group("term").strip(), match.group("text").strip()) if match else None


def split_lead_in(text: str) -> Tuple[str, str]:
    """ "Scope: rest" -> ("Scope", "rest"). A colon with nothing after it, a bracketed placeholder, or a long head is
    not a title."""
    position = text.find(":") if text else -1
    # "[INPUT REQUIRED: fill this in]" is a placeholder, not a titled line
    if position < 0 or text.lstrip().startswith("["):
        return "", text
    head, rest = text[:position], text[position + 1 :].strip()
    if (
        not head.strip()
        or not rest
        or len(head) > cfg.LEAD_IN["max_chars"]
        or len(head.split()) > cfg.LEAD_IN["max_words"]
        or any(char in head for char in cfg.LEAD_IN["stop_chars"])
    ):
        return "", text
    return head.strip(), rest


def resolve_marker(line: str, state: dict) -> Optional[str]:
    """Rewrite a list line as "1." at the depth its marker shape implies, relative to the first marker in the body.
    A new shape lands at most one level under the previous item and keeps that depth, so nothing can skip a level."""
    match = LIST_MARKER.match(line)
    if not match:
        return None
    shape = next(name for name, value in match.groupdict().items() if value)
    level = MARKER_DEPTH[shape]
    depths = state.setdefault("depths", {})  # shape level -> depth
    if level not in depths:
        first = min(depths) if depths else level
        depths[level] = min(max(0, level - first), state.get("prev", -1) + 1)
    depth = depths[level]
    state["prev"] = depth
    return "    " * depth + "1. " + line[match.end() :]


def normalize_lists(lines: List[str]) -> List[str]:
    """Bold-mode body lines made safe for the parser: markers resolved, blanks between items dropped, prose flush
    left."""
    out: List[str] = []
    state: dict = {}
    fence: Optional[List[str]] = None
    for line in lines:
        if fence is not None:
            fence.append(line)
            if FENCE_LINE.match(line):
                out.extend(fence)
                fence = None
            continue
        if FENCE_LINE.match(line):
            fence = [line]
            continue
        if RULE_LINE.match(line):
            continue
        if not line.strip():
            out.append(line)
            continue
        resolved = resolve_marker(line, state)
        if resolved is not None:
            # blank lines between items carry nothing (styles set the
            # spacing) and are the other half of the code-block trigger
            while state.get("in_list") and out and out[-1] == "":
                out.pop()
            state["in_list"] = True
            out.append(resolved)
            continue
        # prose after a blank line ends the list (next marker restarts at depth 0); without a blank it stays with the
        # item
        if out and out[-1] == "":
            state["prev"] = -1
        state["in_list"] = False
        out.append(line.lstrip())
    if fence is not None:
        out.extend(fence)
    return out


def body_nodes(block: Block, doc: DlaiDocument) -> list:
    """Parse one block's lines. Bold-mode input is list-normalised first;
    # input is parsed as written, since its bullets nest by indentation"""
    text_lines = normalize_lists(block.lines) if doc.normalize else block.lines
    nodes = logic.parse_markdown("\n".join(text_lines)).children
    check_code_blocks(nodes, block, doc.findings)
    return nodes


def check_code_blocks(nodes, block: Block, findings: List[str]):
    """A code block in a body is a formatting accident (a fence, or text
    indented four spaces after a blank line). Raise when strict, else note"""

    def walk(items):
        for node in items:
            if node.type in ("code_block", "fence"):
                where = block.source + (node.map[0] + 1 if node.map else 0)
                yield "code block in '%s' near source line %d" % (block.title or "cover", where)
            yield from walk(node.children)

    for problem in walk(nodes):
        if cfg.STRICT_NO_CODE_BLOCKS:
            raise ConversionError(problem)
        if problem not in findings:
            findings.append(problem)


class DlaiWriter(logic.DocxWriter):
    """logic.DocxWriter with the template's rules; only write_heading, write_paragraph and new_paragraph are
    overridden."""

    def __init__(
        self,
        document,
        opts,
        numbering,
        num_id,
        offset=0,
        base_level=-1,
        min_level=0,
        blocks=("enclosure", "encl_h2", "encl_h3"),
        bookmarks=None,
        body_key="prose",
        lead_key="subsection",
        indent_in=None,
    ):
        super().__init__(document, opts)
        self.indent_in = indent_in  # per-enclosure indent step, or None
        self.numbering = numbering  # shared, so numIds stay unique
        self.num_id = num_id
        self.offset = offset
        self.min_level = min_level
        # numbering level of the heading in scope; body text sits one level below it
        self.level = base_level
        self.blocks = blocks
        # Shared list, not a copy: the caller and every writer pop from the
        # same sequence, which keeps bookmarks lined up with the ToC
        self.bookmarks = bookmarks if bookmarks is not None else []
        # DEFINITIONS overrides these so its entries sit flush left and
        # unnumbered, which is the one section that breaks the pattern
        self.body_key, self.lead_key = body_key, lead_key
        # whether the last thing written was ordinary prose, which decides
        # whether a list that follows is subordinate to it
        self.after_prose = False

    def content_level(self, position) -> int:
        depth = max(0, position.list_depth)
        # a list introduced by a paragraph nests under it; a list right after a heading starts at the top level
        if position.list_depth >= 0 and self.after_prose:
            depth += 1
        return self.level + 1 + depth

    def write_heading(self, node, position):
        level = max(1, heading_level(node) - self.offset)
        key = self.blocks[min(level, len(self.blocks)) - 1]
        num_level = max(self.min_level, level - 1)
        paragraph = write(
            self.document, key, heading_text(node), level=num_level, num_id=self.num_id, indent_in=self.indent_in
        )
        if cfg.STYLES[key].numbered:
            self.level = num_level
        if self.bookmarks and level <= cfg.BACK["toc_depth"]:
            helpers.add_bookmark(paragraph, self.bookmarks.pop(0))
        self.after_prose = False
        self.last_paragraph = paragraph
        return paragraph

    def write_paragraph(self, node, position):
        # a hard line break separates words; without this "Overview:" and
        # the text on the next line run together
        spans = logic.inline_spans(logic.find_inline_child(node))
        text = " ".join("".join(" " if span.is_line_break else span.text for span in spans).split())
        if not text:
            return None
        if position.list_depth < 0:
            self.after_prose = True
        lead_in, rest = split_lead_in(text)
        self.last_paragraph = write(
            self.document,
            self.lead_key if lead_in else self.body_key,
            rest,
            level=self.content_level(position),
            lead_in=lead_in,
            num_id=self.num_id,
            indent_in=self.indent_in,
        )
        return self.last_paragraph

    def new_paragraph(self, position, style: Optional[str] = None):
        level = self.content_level(position)
        paragraph = write(self.document, self.body_key, level=level, indent_in=self.indent_in)
        if position.list_depth >= 0:
            helpers.Numbering.apply_to_paragraph(paragraph, self.num_id, level)
        self.last_paragraph = paragraph
        return paragraph


def render_body(
    document,
    opts,
    numbering,
    doc: DlaiDocument,
    block: Block,
    num_id: int = None,
    *,
    base_level=-1,
    min_level=0,
    blocks=("enclosure", "encl_h2", "encl_h3"),
    bookmarks=None,
    body_key="prose",
    lead_key="subsection",
    indent_in=None,
):
    """Write one block's body through DlaiWriter; a fresh cascade is started only if there is something to write."""
    nodes = body_nodes(block, doc)
    if not nodes:
        return
    if num_id is None:
        num_id = new_cascade_list(numbering)
    DlaiWriter(
        document,
        opts,
        numbering,
        num_id,
        offset=doc.offset,
        base_level=base_level,
        min_level=min_level,
        blocks=blocks,
        bookmarks=bookmarks,
        body_key=body_key,
        lead_key=lead_key,
        indent_in=indent_in,
    ).write_blocks(nodes, logic.Position())


def enclosure_spec(title: str):
    spec = cfg.ENCLOSURES.get(normalize(title))
    return spec if spec is not None and spec.enabled else cfg.ENCLOSURES["*"]


# =========================================================================== #
# 3. THE HARDCODED SHELL  -  none of this comes from markdown
# =========================================================================== #
def setup_page(document, doc_title: str):
    """Page geometry, the Normal style, then the header and footer"""
    for section in document.sections:
        section.page_width = Inches(cfg.PAGE["width_in"])
        section.page_height = Inches(cfg.PAGE["height_in"])
        section.top_margin = section.bottom_margin = Inches(cfg.PAGE["margin_in"])
        section.left_margin = section.right_margin = Inches(cfg.PAGE["margin_in"])

    # docDefaults sits underneath every style, so the 1.15 line spacing
    # python-docx ships with has to be cleared there, not on Normal
    helpers.set_document_defaults(document, line_spacing=cfg.LINE_SPACING, space_after_pt=cfg.GAP)
    helpers.define_style(
        document,
        "Normal",
        font=cfg.BODY_FONT,
        size_pt=cfg.BODY_SIZE,
        color=cfg.BODY_COLOR,
        space_after_pt=cfg.GAP,
        line_spacing=cfg.LINE_SPACING,
    )
    create_styles(document)

    section = document.sections[0]
    for area, align, text in (
        (section.header, cfg.PAGE["header_align"], doc_title),
        (section.footer, cfg.PAGE["footer_align"], None),
    ):
        area.is_linked_to_previous = False
        paragraph = area.paragraphs[0]
        paragraph.alignment = align
        if text:
            paragraph.add_run(text)
        else:
            helpers.add_page_number_field(paragraph)


def write_cover(document, numbering, doc_title: str, fields=None, template_type=None):
    """`fields` comes from split_cover. Anything absent falls back to the
    placeholder in config, so the cover still builds from nothing"""
    cover, fields = dict(cfg.COVER), fields or {}
    if template_type:
        cover["doc_type"] = template_type
    for path, _, _ in (cover["seal"], cover["rule"]):
        if not Path(path).is_file():
            raise FileNotFoundError("Cover image missing: %s\nThe assets folder must sit next to " "config.py." % path)
    write_image(document, *cover["seal"])
    write(document, "agency", cover["agency_name"])
    write(document, "doc_type", cover["doc_type"])
    write(document, "cover_line", doc_title)
    effective = " ".join((fields.get("effective") or "").split())
    write(document, "cover_line", cover["effective_pattern"] % effective if effective else cover["effective_text"])
    write_image(document, *cover["rule"], key="agency")

    for index, (text, key) in enumerate(cover["labels"]):
        paragraph = write(document, "cover_label", text)
        value = " ".join((fields.get(key) or "").split()) if key else ""
        if value:
            # a plain run, so the label stays underlined and the value does not
            paragraph.add_run(value)
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(cover["label_space_before_pt"])

    # one reference per line or per semicolon, falling back to placeholders
    supplied = [r.strip() for r in fields.get("references", "").replace(";", "\n").splitlines() if r.strip()]
    entries = supplied or [cover["ref_pattern"] % i for i in range(1, cover["ref_count"] + 1)]

    num_id = new_cascade_list(numbering, cover["ref_cascade"])
    for index, entry in enumerate(entries, start=1):
        paragraph = write(document, "cover_ref", entry, num_id=num_id)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(cover["ref_indent_in"])
        paragraph_format.first_line_indent = Inches(-cfg.INDENT_STEP_IN)
        if index == len(entries):
            paragraph_format.space_after = Pt(cfg.GAP)


def write_signature_block(document, titles: List[str], lines=None, glossary: bool = True):
    """Signature lines, then the Enclosure(s) list exactly as the enclosures were found, Glossary last and
    unnumbered."""
    back = cfg.BACK
    for line in lines or back["sig_lines"]:
        write(document, "signature", line)
    write(document, "encl_label", back["encl_label"])
    for index, title in enumerate(titles, start=1):
        write(document, "encl_item", back["encl_pattern"] % (index, title))
    if glossary:
        last = write(document, "encl_item", back["encl_last"])
        last.paragraph_format.first_line_indent = Inches(0)


class TocEntry(NamedTuple):
    level: int
    text: str
    bookmark: str


class BackPart(NamedTuple):
    key: str  # heading style: enclosure, appendices or glossary_part
    level: int  # table of contents level
    title: str
    block: Optional[Block]
    writer: object  # a glossary writer, or None for flat prose


def back_matter_plan(doc: DlaiDocument) -> List[BackPart]:
    """Everything after the enclosures, in print order; shared by plan_toc and write_back_matter so bookmarks line
    up."""
    conf = cfg.BACK
    parts: List[BackPart] = []
    if doc.appendices is not None:
        parts.append(BackPart("appendices", 1, conf["appendices_title"], doc.appendices, None))
    if doc.glossary is not None:
        parts.append(BackPart("enclosure", 1, conf["encl_last"], doc.glossary.preamble, None))
        if doc.glossary.abbreviations is not None:
            parts.append(
                BackPart(
                    "glossary_part",
                    2,
                    doc.glossary.abbreviations.title,
                    doc.glossary.abbreviations,
                    write_glossary_columns,
                )
            )
        if doc.glossary.definitions is not None:
            parts.append(
                BackPart(
                    "glossary_part",
                    2,
                    doc.glossary.definitions.title,
                    doc.glossary.definitions,
                    write_glossary_definitions,
                )
            )
        parts += [BackPart("glossary_part", 2, part.title, part, None) for part in doc.glossary.other]
    else:
        parts.append(BackPart("enclosure", 1, conf["encl_last"], None, None))
    for name, block in (("TABLES", doc.tables), ("FIGURES", doc.figures)):
        if block is not None:
            parts.append(BackPart("enclosure", 1, name, block, None))
    return parts


def plan_toc(doc: DlaiDocument) -> List[TocEntry]:
    """Every ToC row with its bookmark name, in document order, before
    anything is written. Enclosure sub-headings count down to toc_depth"""
    entries: List[TocEntry] = []

    def add(level, text):
        entries.append(TocEntry(level, text, "DLAIREF%d" % (len(entries) + 1)))

    for index, block in enumerate(doc.enclosures, start=1):
        add(1, cfg.BACK["encl_display"] % (index, block.title))
        for node in body_nodes(block, doc):
            if node.type == "heading":
                level = max(1, heading_level(node) - doc.offset)
                if level <= cfg.BACK["toc_depth"]:
                    add(level, heading_text(node))
    for part in back_matter_plan(doc):
        add(part.level, part.title)
    return entries


def write_table_of_contents(document, entries: List[TocEntry]):
    """A real Word TOC field mapped to our styles by name; Word computes pages, indents and leaders on open."""
    write(document, "toc_title", cfg.BACK["toc_title"])

    mapping = ",".join("%s,%d" % (cfg.STYLES[key].name, level) for key, level in cfg.BACK["toc_levels"])
    paragraph = document.add_paragraph()
    helpers.add_field(paragraph, r' TOC \h \z \t "%s" ' % mapping, placeholder=cfg.BACK["toc_placeholder"])


def write_glossary_columns(document, body, opts, numbering):
    """PART I rows: term, tab, meaning; sorted, grouped by initial letter; non-row lines print flush left after."""
    columns = cfg.GLOSSARY_COLUMNS
    lines, blocks = [], []
    for node in body:
        (lines.extend(paragraph_lines(node)) if node.type == "paragraph" else blocks.append(node))

    rows, leftovers = [], []
    for line in lines:
        entry = glossary_entry(line)
        (rows if entry else leftovers).append(entry or line)
    rows.sort(key=lambda row: row[0].upper())

    previous = None
    for term, meaning in rows:
        initial = term[:1].upper()
        paragraph = write(document, "glossary_entry")
        fmt = paragraph.paragraph_format
        fmt.left_indent = Inches(columns["column_in"])
        fmt.first_line_indent = Inches(-columns["column_in"])
        fmt.tab_stops.add_tab_stop(Inches(columns["column_in"]))
        if previous is not None and initial != previous:
            fmt.space_before = Pt(columns["group_gap_pt"])
        paragraph.add_run(term)
        paragraph.add_run("\t")
        paragraph.add_run(meaning)
        previous = initial

    for line in leftovers:
        write(document, "flat", line)
    if blocks:
        DlaiWriter(
            document,
            opts,
            numbering,
            new_cascade_list(numbering),
            base_level=-1,
            min_level=0,
            blocks=("subsection",),
            body_key="flat",
            lead_key="definition",
        ).write_blocks(blocks, logic.Position())


def write_glossary_definitions(document, body, opts, numbering):
    """PART II: one definition per paragraph, "Term.  text", the term underlined, sorted."""
    lines, blocks = [], []
    for node in body:
        (lines.extend(paragraph_lines(node)) if node.type == "paragraph" else blocks.append(node))

    entries, leftovers = [], []
    for line in lines:
        entry = glossary_entry(line)
        (entries if entry else leftovers).append(entry or line)
    if cfg.GLOSSARY_DEFINITIONS["sort"]:
        entries.sort(key=lambda entry: entry[0].upper())

    for term, text in entries:
        write(document, "glossary_definition", text, lead_in=term)
    for line in leftovers:
        write(document, "flat", line)
    if blocks:
        DlaiWriter(
            document,
            opts,
            numbering,
            new_cascade_list(numbering),
            base_level=-1,
            min_level=0,
            blocks=("subsection",),
            body_key="flat",
            lead_key="glossary_definition",
        ).write_blocks(blocks, logic.Position())


def write_back_matter(document, bookmarks: List[str], doc: DlaiDocument, opts, numbering):
    """Appendices, the glossary with its parts, then tables and figures, each only when the author supplied it."""
    for part in back_matter_plan(doc):
        paragraph = write(document, part.key, part.title)
        helpers.add_bookmark(paragraph, bookmarks.pop(0))
        if part.block is None:
            continue
        if part.writer is not None:
            part.writer(document, body_nodes(part.block, doc), opts, numbering)
        else:
            render_body(
                document,
                opts,
                numbering,
                doc,
                part.block,
                base_level=-1,
                min_level=0,
                blocks=("subsection",),
                body_key="flat",
                lead_key="definition",
            )


# =========================================================================== #
# 4. ASSEMBLY
# =========================================================================== #
def md_to_docx(doc: DlaiDocument, doc_title: str = "PLACEHOLDER TITLE", template_name=None):
    """Render a DlaiDocument (from boundaries.build_document or the section
    agent) as a .docx. Returns (bytes, findings)."""
    findings = doc.findings
    opts = helpers.DocxOptions(body_font=cfg.BODY_FONT, body_size_pt=cfg.BODY_SIZE)
    document = WordDocument()
    numbering = helpers.Numbering(document)

    # 1. every ToC row and bookmark name, before anything is written
    entries = plan_toc(doc)
    bookmarks = [entry.bookmark for entry in entries]

    # 2. page setup, the Word styles, then the cover
    setup_page(document, doc_title)
    write_cover(document, numbering, doc_title, doc.cover, template_name)

    # 3. the required sections, in config order
    write_sections(document, opts, numbering, doc)

    # 4. signature block and table of contents
    write_signature_block(document, [b.title for b in doc.enclosures], doc.signature)
    write_table_of_contents(document, entries)

    # 5. the enclosures, then whatever back matter the author supplied
    write_enclosures(document, opts, numbering, doc, bookmarks)
    write_back_matter(document, bookmarks, doc, opts, numbering)

    # 6. tell Word to refresh page numbers on open, then serialise to bytes
    helpers.update_fields_on_open(document)
    return helpers.document_to_bytes(document), findings


def build_dlai(
    markdown_input: str,
    sections_input: str = "",
    doc_title: str = "PLACEHOLDER TITLE",
    template_name=None,
    provider: str = None,
):
    """Markdown text in, (bytes, findings) out: build_document then md_to_docx. The older two-string call still
    works."""
    text = "\n".join(part for part in (sections_input, markdown_input) if part)
    return md_to_docx(boundaries.build_document(text, provider), doc_title, template_name)


def write_sections(document, opts, numbering, doc: DlaiDocument):
    """The numbered sections in config order; a missing required section prints the placeholder, an optional one is
    skipped."""
    num_id = new_cascade_list(numbering)
    first = True
    for name, spec in cfg.SECTIONS.items():
        if name == cfg.SIGNATURE_SECTION:
            continue
        block = doc.sections[name]
        if block is not None and not body_nodes(block, doc):
            doc.findings.append("EMPTY section: %s%s" % (name, "" if spec.optional else " (placeholder written)"))
            block = None
        if block is None and spec.optional:
            continue  # left out entirely; the numbering closes up
        if first and cfg.PAGE["break_after_cover"]:
            new_page(document)
        first = False
        write(document, "section", name, num_id=num_id)
        if block is None:
            write(document, "prose", cfg.MISSING_PLACEHOLDER, level=1, num_id=num_id)
            continue
        render_body(
            document,
            opts,
            numbering,
            doc,
            block,
            num_id,
            base_level=0,
            min_level=1,
            blocks=("subsection",),
            body_key=spec.body_key,
            lead_key=spec.lead_key,
        )


def write_enclosures(document, opts, numbering, doc: DlaiDocument, bookmarks: List[str]):
    """A fresh numId per enclosure, which restarts its numbering at 1, with
    that enclosure's cascade and indent from config.ENCLOSURES"""
    for index, block in enumerate(doc.enclosures, start=1):
        spec = enclosure_spec(block.title)
        num_id = new_cascade_list(numbering, spec.cascade)
        heading = write(document, "enclosure", cfg.BACK["encl_display"] % (index, block.title))
        helpers.add_bookmark(heading, bookmarks.pop(0))
        render_body(document, opts, numbering, doc, block, num_id, bookmarks=bookmarks, indent_in=spec.indent_in)


def split_input(markdown_text: str) -> Tuple[str, str]:
    """Split one combined string into (enclosures, sections) by heading name. Kept for older callers; build_dlai rejoins
    them."""
    sections, enclosures = [], []
    # anything above the first heading is the cover block, so it starts on the sections side
    target, depth = sections, 99

    for line in (markdown_text or "").splitlines():
        match = HEADING_LINE.match(line)
        if match and len(match.group(1)) <= depth:
            name = SECTION_ALIASES.get(normalize(match.group(2)))
            target = sections if name else enclosures
            depth = len(match.group(1)) if name else 99
        target.append(line)

    return ("\n".join(enclosures).strip() + "\n", "\n".join(sections).strip() + "\n")
