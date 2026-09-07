from __future__ import annotations
import html as html_module
import re
from dataclasses import dataclass, field, replace
from typing import List, Optional, Tuple
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode
import helpers
from helpers import DocxOptions

TEXT_ALIGNMENTS = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,}
LIST_INDENT_INCHES = 0.25       # per nesting level, and the hanging indent size
QUOTE_INDENT_INCHES = 0.3       # per blockquote level
SPACE_AFTER_BLOCK = Pt(8)       # normal gap once a list or quote block ends
SPACE_BETWEEN_LIST_ITEMS = Pt(2)

FRONT_MATTER_PATTERN = re.compile(r"\A(?:---|\+\+\+)\r?\n.*?\r?\n(?:---|\+\+\+)\s*?\r?\n", re.S)
ILLEGAL_XML_CHARS = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff\ufffe\uffff]")
HTML_TAG_PATTERN = re.compile(r"<[^>]*>")
LINE_BREAK_TAG_PATTERN = re.compile(r"<\s*br\s*/?\s*>", re.I)
TASK_CHECKBOX_PATTERN = re.compile(r"^\[([ xX])\]\s+")

def clean_source(text, strip_front_matter: bool = True) -> str:
    text = text if isinstance(text, str) else str(text or "")
    text = text.removeprefix("\ufeff").replace("\r\n", "\n").replace("\r", "\n")
    text = ILLEGAL_XML_CHARS.sub("", text)
    if strip_front_matter:
        text = FRONT_MATTER_PATTERN.sub("", text, count=1)
    return text

def strip_html_tags(raw_html: str) -> str:
    return html_module.unescape(HTML_TAG_PATTERN.sub("", raw_html))

MARKDOWN_PARSER = MarkdownIt("commonmark").enable(["table", "strikethrough"])

def parse_markdown(text: str) -> SyntaxTreeNode:
    return SyntaxTreeNode(MARKDOWN_PARSER.parse(text))

def find_inline_child(node) -> Optional[SyntaxTreeNode]:
    return next((child for child in node.children if child.type == "inline"), None)

@dataclass(frozen=True)
class Span:
    """One run of text with one set of formatting, or a line break."""
    text: str = ""
    bold: bool = False
    italic: bool = False
    strikethrough: bool = False
    is_code: bool = False
    link_url: str = ""          # empty when this span is not inside a link
    is_line_break: bool = False

def inline_spans(node) -> List[Span]:
    collected: List[Span] = []
    if node is not None:
        collect_spans(node, {}, collected)
    return collected

def collect_spans(node, active_formatting: dict, collected: List[Span]):
    FORMATTING_NODES = {"strong": "bold", "em": "italic", "s": "strikethrough"}
    for child in node.children:
        node_type = child.type
        if node_type == "text":
            if child.content:
                collected.append(Span(text=child.content, **active_formatting))
        elif node_type == "code_inline":
            collected.append(Span(text=child.content, is_code=True, **active_formatting))
        elif node_type in FORMATTING_NODES:
            collect_spans(child,{**active_formatting, FORMATTING_NODES[node_type]: True},collected)
        elif node_type == "link":
            collect_spans(child,{**active_formatting, "link_url": child.attrs.get("href", "")},collected)
        elif node_type == "softbreak":
            collected.append(Span(text=" ", **active_formatting))
        elif node_type == "hardbreak":
            collected.append(Span(is_line_break=True))
        elif node_type == "html_inline":
            if LINE_BREAK_TAG_PATTERN.fullmatch(child.content.strip()):
                collected.append(Span(is_line_break=True))
            else:
                text = strip_html_tags(child.content)
                if text:
                    collected.append(Span(text=text, **active_formatting))
        else:
            # Unknown wrapper (and images): keep the contents, drop the wrapper.
            # An image node's only child is its alt text, so ![alt](src) degrades
            # to that alt text as plain text. This is why images need no handling.
            collect_spans(child, active_formatting, collected)

def as_task_item(spans: List[Span]) -> List[Span]:
    if not spans or not spans[0].text:
        return spans
    match = TASK_CHECKBOX_PATTERN.match(spans[0].text)
    if not match:
        return spans
    checkbox = "\u2611\ufe0e " if match.group(1).lower() == "x" else "\u2610\ufe0e "
    first_span = replace(spans[0], text=checkbox + spans[0].text[match.end():])
    return [first_span] + spans[1:]

def spans_text(spans: List[Span]) -> str:
    return "".join(span.text for span in spans)

@dataclass
class Cell:
    spans: List[Span] = field(default_factory=list)
    alignment: Optional[str] = None

def table_grid(node) -> Tuple[List[List[Cell]], int]:
    header_rows: List[List[Cell]] = []
    body_rows: List[List[Cell]] = []
    for section in node.children:
        (header_rows if section.type == "thead" else body_rows).extend([read_cell(cell) for cell in row.children] for row in section.children)
    all_rows = [row for row in header_rows + body_rows if row]
    return all_rows, len(header_rows)

def read_cell(node) -> Cell:
    style_attribute = str(node.attrs.get("style", ""))
    alignment = next((name for name in ("center", "right", "left") if name in style_attribute), None)
    return Cell(spans=inline_spans(find_inline_child(node)), alignment=alignment)

NODES = {
    "heading": "write_heading", "paragraph": "write_paragraph",
    "bullet_list": "write_list", "ordered_list": "write_list",
    "list_item": "write_list_item", "blockquote": "write_blockquote",
    "fence": "write_code_block", "code_block": "write_code_block",
    "hr": "write_horizontal_rule", "html_block": "write_html_block",
    "table": "write_table",
}

@dataclass(frozen=True)
class Position:
    quote_depth: int = 0
    list_depth: int = -1
    list_num_id: int = 0
    is_first_line_of_item: bool = False

class DocxWriter:
    def __init__(self, document, opts: DocxOptions):
        self.document = document
        self.opts = opts
        self.numbering = helpers.Numbering(document)
        self.last_paragraph = None

    def write_document(self, tree_root):
        self.write_blocks(tree_root.children, Position())

    def write_blocks(self, nodes, position: Position):
        for node in nodes:
            handler_name = NODES.get(node.type)
            if handler_name is not None:
                getattr(self, handler_name)(node, position)
            elif node.children:
                self.write_blocks(node.children, position)

    def write_heading(self, node, position: Position):
        heading_level = min(int(node.tag[1:] or 1), 6)
        paragraph = self.new_paragraph(position, style="Heading %d" % heading_level)
        self.write_spans(paragraph, inline_spans(find_inline_child(node)))

    def write_paragraph(self, node, position: Position):
        spans = inline_spans(find_inline_child(node))
        if position.is_first_line_of_item:
            spans = as_task_item(spans)
        paragraph = self.new_paragraph(position)
        self.write_spans(paragraph, spans)
        is_empty = (not paragraph.runs and not paragraph._p.findall(qn("w:hyperlink")))
        if position.list_depth < 0 and is_empty:
            paragraph._p.getparent().remove(paragraph._p)

    def write_list(self, node, position: Position):
        is_ordered = node.type == "ordered_list"
        try:
            start_number = int(node.attrs.get("start", 1) or 1)
        except (TypeError, ValueError):
            start_number = 1
        inside_list = replace(position, list_depth=position.list_depth + 1,
            list_num_id=self.numbering.create_list(is_ordered, start_number),
            is_first_line_of_item=False)
        self.write_blocks(node.children, inside_list)
        if position.list_depth < 0:
            self.restore_spacing_after_block()

    def write_list_item(self, node, position: Position):
        for block_index, child in enumerate(node.children):
            self.write_blocks([child], replace(position, is_first_line_of_item=block_index == 0))

    def write_blockquote(self, node, position: Position):
        self.write_blocks(node.children,
            replace(position, quote_depth=position.quote_depth + 1, is_first_line_of_item=False))
        if position.quote_depth == 0:
            self.restore_spacing_after_block()

    def write_code_block(self, node, position: Position):
        paragraph = self.new_paragraph(position)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        list_indent = (LIST_INDENT_INCHES * (position.list_depth + 1)
                       if position.list_depth >= 0 else 0)
        paragraph_format.left_indent = Inches(0.15 + list_indent + QUOTE_INDENT_INCHES * position.quote_depth)
        paragraph_format.right_indent = Inches(0.1)
        helpers.keep_paragraph_together(paragraph)
        helpers.shade_paragraph(paragraph, self.opts.code_fill)
        helpers.set_paragraph_borders(paragraph,
            {side: {"sz": 4, "color": "D9D9D9", "space": 6} for side in ("top", "left", "bottom", "right")})
        for line_index, line in enumerate(node.content.rstrip("\n").split("\n")):
            if line_index:
                self.apply_code_font(paragraph.add_run()).add_break()
            self.apply_code_font(paragraph.add_run(line.replace("\t", "    ")))

    def write_horizontal_rule(self, node, position: Position):
        paragraph = self.new_paragraph(position)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        helpers.set_paragraph_borders(paragraph, {"bottom": {"sz": 6, "color": "A6A6A6", "space": 1}})

    def write_html_block(self, node, position: Position):
        text = strip_html_tags(LINE_BREAK_TAG_PATTERN.sub("\n", node.content)).strip()
        for raw_line in text.split("\n"):
            if (line := raw_line.strip()):
                self.write_spans(self.new_paragraph(position), [Span(text=line)])

    def write_table(self, node, position: Position):
        rows, header_row_count = table_grid(node)
        if not rows:
            return
        column_count = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=column_count)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        column_widths = self.calculate_column_widths(rows, column_count)
        for row_index, cells in enumerate(rows):
            is_header_row = row_index < header_row_count
            table_row = table.rows[row_index]
            if is_header_row:
                helpers.mark_as_repeating_header(table_row)
            for column_index in range(column_count):
                cell = table_row.cells[column_index]
                cell.width = column_widths[column_index]
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                # markdown-it pads and truncates every row to the header width,
                # so cells[column_index] is always present, no bounds check needed
                cell_content = cells[column_index]
                paragraph.alignment = TEXT_ALIGNMENTS.get(cell_content.alignment)
                self.write_spans(paragraph, cell_content.spans)
                if is_header_row:
                    helpers.shade_table_cell(cell, self.opts.table_header_fill)
                    link_runs = [run for link in paragraph.hyperlinks for run in link.runs]
                    for run in list(paragraph.runs) + link_runs:
                        run.bold = True
        self.document.add_paragraph()

    def new_paragraph(self, position: Position, style: Optional[str] = None):
        try:
            paragraph = self.document.add_paragraph(style=style)
        except KeyError:
            paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        quote_indent = QUOTE_INDENT_INCHES * position.quote_depth
        if position.quote_depth:
            helpers.set_paragraph_borders(paragraph,
                {"left": {"sz": 18, "color": self.opts.quote_bar, "space": 10}})
        if position.list_depth >= 0:
            paragraph_format.left_indent = Inches(LIST_INDENT_INCHES * (position.list_depth + 1) + quote_indent)
            if position.is_first_line_of_item:
                self.numbering.apply_to_paragraph(paragraph, position.list_num_id, position.list_depth)
                paragraph_format.first_line_indent = Inches(-LIST_INDENT_INCHES)
            paragraph_format.space_after = SPACE_BETWEEN_LIST_ITEMS
        elif quote_indent:
            paragraph_format.left_indent = Inches(quote_indent)
        self.last_paragraph = paragraph
        return paragraph

    def restore_spacing_after_block(self):
        if self.last_paragraph is not None:
            self.last_paragraph.paragraph_format.space_after = SPACE_AFTER_BLOCK

    def write_spans(self, paragraph, spans: List[Span]):
        """Write each span as a run, wrapping any linked run in a hyperlink.

        A markdown link that spans several runs (e.g. [a **b** c](url)) becomes
        several adjacent hyperlinks to the same target rather than one wrapping
        all of them. python-docx points them at a single shared relationship, so
        the result is visually identical and clickable, just simpler to build.
        """
        for span in spans:
            run = self.write_span(paragraph, span)
            if run is not None and span.link_url:
                self.wrap_run_in_hyperlink(paragraph, span.link_url, run)

    def write_span(self, paragraph, span: Span):
        if span.is_line_break:
            paragraph.add_run().add_break()
            return None
        run = paragraph.add_run(span.text)
        run.bold = span.bold or None
        run.italic = span.italic or None
        if span.strikethrough:
            run.font.strike = True
        if span.is_code:
            self.apply_code_font(run)
            helpers.shade_run(run, self.opts.code_fill)
        elif self.opts.body_font:
            run.font.name = self.opts.body_font
        if span.link_url:
            run.font.color.rgb = RGBColor.from_string(self.opts.link_color)
            run.font.underline = True
        return run

    def apply_code_font(self, run):
        helpers.apply_monospace_font(run, self.opts.code_font, self.opts.code_size_pt)
        return run

    def wrap_run_in_hyperlink(self, paragraph, url: str, run):
        """Make one already-written run clickable.

        A Word hyperlink is not a run property: it is a <w:hyperlink> element
        that wraps the run and points, by id, at a relationship in a separate
        part of the zip. python-docx has no API for this, so we register the
        relationship, build the wrapper, and move the run inside it. An unusable
        target is swallowed, leaving the styled-but-not-clickable text.
        """
        try:
            relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
        except Exception:
            return
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run._element.addprevious(hyperlink)     # drop the wrapper where the run was
        hyperlink.append(run._element)          # then move the run into it

    def calculate_column_widths(self, rows, column_count) -> List[Emu]:
        section = self.document.sections[0]
        printable_width = (section.page_width - section.left_margin - section.right_margin)
        weights = []
        for column_index in range(column_count):
            longest_text = max((len(spans_text(row[column_index].spans)) for row in rows if column_index < len(row)), default=1)
            weights.append(min(max(longest_text, 6), 60))
        total_weight = sum(weights)
        minimum_width = int(Inches(0.6))
        widths = [max(int(printable_width * weight / total_weight), minimum_width) for weight in weights]
        shrink_factor = printable_width / sum(widths)
        if shrink_factor < 1:
            widths = [int(width * shrink_factor) for width in widths]
        return [Emu(width) for width in widths]
